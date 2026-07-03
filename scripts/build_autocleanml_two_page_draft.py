from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("AutoCleanML_Committee_Refresher.docx")


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.86)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    configure_styles(doc)
    add_footer(doc)

    add_title(doc)
    add_overview(doc)
    add_research_questions(doc)
    add_basic_framework(doc)
    add_evaluation_plan(doc)
    add_committee_takeaway(doc)

    doc.save(OUTPUT)
    print(OUTPUT.resolve())


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.8)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Heading 1", 14, "2E74B5", 10, 4),
        ("Heading 2", 12, "1F4D78", 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("AutoCleanML committee refresher")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("AutoCleanML")
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(6)
    run = subtitle.add_run(
        "An automated data-centric framework for improving data quality in "
        "Spark-based machine learning workflows"
    )
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(80, 80, 80)


def add_overview(doc: Document) -> None:
    doc.add_heading("Purpose and Motivation", level=1)
    add_body(
        doc,
        "Machine learning systems depend heavily on the quality of the data "
        "used to train them. In big-data environments, this data often contains "
        "missing values, duplicate records, outliers, schema changes, skewed "
        "keys, and label-related issues. Today, these problems are usually "
        "handled through one-off cleaning scripts written for a specific "
        "dataset. That makes the cleaning process hard to repeat, hard to "
        "compare across projects, and difficult to maintain as the data "
        "changes."
    )
    add_body(
        doc,
        "AutoCleanML is proposed as a reusable Spark-based framework that "
        "treats data cleaning as a data-centric engineering process. Instead "
        "of manually writing a new cleaning script for each dataset, the user "
        "selects general cleaning policies, and the framework profiles the "
        "data, applies repairs, and records what changed. The goal is not to "
        "claim that every model will always improve, but to make data-quality "
        "improvement more systematic, measurable, and reproducible."
    )


def add_research_questions(doc: Document) -> None:
    doc.add_heading("Research Questions", level=1)
    add_labeled_item(
        doc,
        "RQ1 - Detection:",
        "How can common data-quality issues in large-scale Spark datasets be automatically profiled and quantified?",
    )
    add_labeled_item(
        doc,
        "RQ2 - Repair:",
        "Can configurable, rule-driven policies repair these issues while reducing the need for manual dataset-specific cleaning code?",
    )
    add_labeled_item(
        doc,
        "RQ3 - Evaluation:",
        "How do automated cleaning operations affect data quality, downstream ML performance, and operational effort?",
    )


def add_basic_framework(doc: Document) -> None:
    doc.add_heading("What AutoCleanML Does", level=1)
    add_body(
        doc,
        "At a high level, AutoCleanML follows a simple cycle: profile the raw "
        "data, repair the detected issues using reusable policies, and then "
        "evaluate the result. The profiling step identifies where the dataset "
        "has quality problems. The repair step applies selected policies such "
        "as filling missing values, removing duplicates, capping outliers, "
        "aligning schemas, or handling missing labels. The evaluation step "
        "compares the raw and cleaned datasets so the effect of cleaning is "
        "visible rather than assumed."
    )
    add_body(
        doc,
        "The important idea is that the policies are configurable but reusable. "
        "For example, choosing a missing-value strategy or duplicate-handling "
        "strategy is still a human decision, but it is not the same as writing "
        "custom cleaning logic for every dataset. AutoCleanML separates the "
        "cleaning decision from the low-level Spark implementation, making the "
        "same workflow usable across multiple datasets."
    )


def add_evaluation_plan(doc: Document) -> None:
    doc.add_heading("How the Thesis Evaluates It", level=1)
    add_body(
        doc,
        "The thesis evaluates AutoCleanML from three angles. First, it checks "
        "whether the framework can detect and reduce data-quality problems "
        "such as missingness, duplicates, outliers, skew, and schema issues. "
        "Second, it compares machine-learning results before and after "
        "cleaning, using classification and regression metrics where "
        "appropriate. Third, it studies operational effort by comparing the "
        "reusable AutoCleanML workflow with manual dataset-specific cleaning "
        "scripts."
    )
    add_body(
        doc,
        "The planned evidence includes both synthetic and real datasets. "
        "Synthetic datasets provide controlled quality problems, while real "
        "datasets show how the framework behaves on common benchmark-style "
        "data. The stronger claim is that cleaning becomes repeatable, "
        "auditable, and less dependent on hand-written scripts."
    )


def add_committee_takeaway(doc: Document) -> None:
    doc.add_heading("Main Contribution", level=1)
    add_body(
        doc,
        "The main contribution of AutoCleanML is a practical framework for "
        "bringing data-centric quality improvement into Spark-based ML "
        "pipelines. It connects detection, repair, and evaluation into one "
        "repeatable process. This supports the thesis argument that data "
        "cleaning should not remain an informal preprocessing step, but should "
        "be treated as an explicit, measurable part of the ML workflow."
    )
    add_body(
        doc,
        "For committee members, the key point is that AutoCleanML is not "
        "intended to replace all human judgment. The user still chooses "
        "reasonable cleaning policies based on the dataset and task. The value "
        "of the framework is that once those policies are selected, the same "
        "Spark-native process can be applied consistently, measured, and reused "
        "without writing a new custom cleaning script each time."
    )
    doc.add_heading("Current Status", level=1)
    add_body(
        doc,
        "The current implementation aligns with the proposal at a functional "
        "level. It supports automated profiling, policy-based repair, raw-versus-"
        "cleaned comparison, optional ML evaluation, experiment artifacts, and "
        "a PDSA-style quality loop. The remaining thesis work is mainly to "
        "present the experiments clearly, explain the limits of the approach, "
        "and frame the results around reproducibility and operational effort."
    )
    add_labeled_item(
        doc,
        "Best framing:",
        "AutoCleanML provides a reusable Spark-native data-quality layer for profiling, repairing, and evaluating common tabular data issues before ML training.",
    )
    add_labeled_item(
        doc,
        "Careful wording:",
        "The thesis should claim reduced manual effort and improved measurable data quality, while treating runtime and ML gains as dataset-dependent outcomes.",
    )


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.font.size = Pt(10.8)


def add_labeled_item(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.06
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    label_run = p.add_run(label + " ")
    label_run.bold = True
    label_run.font.size = Pt(10.6)
    text_run = p.add_run(text)
    text_run.font.size = Pt(10.6)


if __name__ == "__main__":
    main()
